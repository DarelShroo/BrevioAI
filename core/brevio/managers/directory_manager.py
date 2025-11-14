import asyncio
import logging
from collections.abc import AsyncGenerator
from typing import Generator, Optional

from core.shared.models.history_token_model import HistoryTokenModel

try:
    import aiofiles
except ImportError:
    raise ImportError("aiofiles is not installed")
import markdown
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.brevio.constants.directory_messages import DirectoryMessages
from core.brevio.constants.summary_messages import SummaryMessages
from core.brevio.models.config_model import ConfigModel
from core.brevio.models.response_model import FolderResponse

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
logger.propagate = False
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)


class DirectoryManager:
    def __init__(self, config: Optional[ConfigModel] = None) -> None:
        logger.info("DirectoryManager initialized")
        self._config: Optional[ConfigModel] = config if config else None

    async def createFolder(self, path: Optional[str] = None) -> FolderResponse:
        if path is None:
            raise ValueError("El path no puede ser None.")
        try:
            await asyncio.to_thread(__import__("os").makedirs, path, exist_ok=True)
            exists = await asyncio.to_thread(__import__("os").path.exists, path)
            if exists:
                return FolderResponse(
                    success=True,
                    message=DirectoryMessages.SUCCESS_FOLDER_CREATED.format(path),
                )
            raise RuntimeError("La carpeta no se creó, la ruta puede ser incorrecta.")
        except OSError as e:
            raise OSError(f"Error al crear la carpeta: {e}")
        except Exception as e:
            raise RuntimeError(f"Error inesperado al crear la carpeta: {str(e)}") from e

    async def deleteFolder(self) -> FolderResponse:
        folder: str = self._config.dest_folder if self._config else ""
        try:
            exists = await asyncio.to_thread(__import__("os").path.exists, folder)
            if not exists:
                raise FileNotFoundError(
                    DirectoryMessages.ERROR_FOLDER_NOT_FOUND.format(folder)
                )
            await asyncio.to_thread(__import__("shutil").rmtree, folder)
            return FolderResponse(
                success=True,
                message=DirectoryMessages.SUCCES_DELETION.format(folder),
            )
        except OSError as e:
            raise OSError(f"Error al eliminar la carpeta: {str(e)}") from e
        except Exception as e:
            raise RuntimeError(
                f"Error inesperado al eliminar la carpeta: {str(e)}"
            ) from e

    async def deleteFile(self, file_path: str) -> FolderResponse:
        try:
            exists = await asyncio.to_thread(__import__("os").path.exists, file_path)
            if not exists:
                raise FileNotFoundError(
                    DirectoryMessages.ERROR_FILE_NOT_FOUND.format(file_path)
                )
            await asyncio.to_thread(__import__("os").remove, file_path)
            return FolderResponse(
                success=True,
                message=DirectoryMessages.SUCCESS_FILE_DELETED.format(file_path),
            )
        except PermissionError:
            raise PermissionError(
                DirectoryMessages.ERROR_PERMISSION_DENIED.format(file_path)
            )
        except OSError as e:
            raise OSError(f"Error al eliminar el archivo: {str(e)}") from e
        except Exception as e:
            raise RuntimeError(
                f"Error inesperado al eliminar el archivo: {str(e)}"
            ) from e

    async def validate_paths(self, transcription_path: str) -> None:
        exists = await asyncio.to_thread(
            __import__("os").path.exists, transcription_path
        )
        if not exists:
            raise FileNotFoundError(
                SummaryMessages.ERROR_READING_TRANSCRIPTION.format(transcription_path)
            )

    async def read_transcription(self, transcription_path: str) -> str:
        try:
            async with aiofiles.open(transcription_path, "r", encoding="utf-8") as f:
                transcription = await f.read()

            if not transcription.strip():
                raise ValueError(SummaryMessages.ERROR_EMPTY_TRANSCRIPTION)

            logger.info(f"Transcripción leída: {len(transcription)} caracteres")
            return transcription
        except FileNotFoundError:
            raise FileNotFoundError(
                f"El archivo de transcripción no se encontró en la ruta: {transcription_path}"
            )
        except ValueError as e:
            raise ValueError(f"Transcripción vacía: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"Error al leer la transcripción: {str(e)}") from e

    async def read_pdf_in_pieces(
        self,
        path: str,
        history_token_model: Optional["HistoryTokenModel"] = None,
        pieces_for_page: int = 1,
    ) -> AsyncGenerator[str, None]:
        def blocking_read() -> list[str]:
            import re

            import pypdf
            import tiktoken

            lector = pypdf.PdfReader(path)
            total_paginas = len(lector.pages)
            logger.info(f"PDF tiene {total_paginas} páginas")
            encoder = tiktoken.encoding_for_model("gpt-4")

            fragments = []
            for i, pagina in enumerate(lector.pages):
                texto = pagina.extract_text() or ""
                texto_limpio = re.sub(r"\n+", "\n", texto.strip())
                if texto_limpio:
                    if history_token_model:
                        history_token_model.num_chars_file += len(texto_limpio)
                        history_token_model.num_tokens_file += len(
                            encoder.encode(texto_limpio)
                        )
                    fragments.append(texto_limpio)
            return fragments

        fragments = await asyncio.to_thread(blocking_read)
        for fragment in fragments:
            yield fragment

    async def read_pdf(
        self, pdf_path: str, history_token_model: Optional["HistoryTokenModel"]
    ) -> list[str]:
        fragments = []
        async for fragment in self.read_pdf_in_pieces(pdf_path, history_token_model):
            fragments.append(fragment)

        if not fragments:
            raise ValueError("El PDF está vacío o no se pudo extraer texto")

        logger.info(
            f"PDF leído: {len(fragments)} fragmentos, Total caracteres: {sum(len(f) for f in fragments)}"
        )
        return fragments

    async def read_docx(self, docx_path: str) -> str:
        def blocking_read() -> str:
            doc = Document(docx_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            total_paragraphs = len(paragraphs)
            logger.info(f"DOCX tiene {total_paragraphs} párrafos")
            return "\n".join(paragraphs)

        return await asyncio.to_thread(blocking_read)

    async def write_summary(self, summary: str, summary_path: str) -> None:
        try:
            async with aiofiles.open(summary_path, "a", encoding="utf-8") as f:
                await f.write(summary + "\n")
            logger.info(f"Resumen escrito en {summary_path}")
        except Exception as e:
            raise RuntimeError(
                f"Error al escribir el resumen en {summary_path}: {str(e)}"
            ) from e

    def _add_toc(self, doc: DocumentType) -> None:
        logger.debug("Adding table of contents to DOCX")
        doc.add_heading("Índice", level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char)
        instr_text = OxmlElement("w:instrText")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr_text)
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char)
        doc.add_paragraph(
            "Nota: Haga clic derecho en el índice y seleccione 'Actualizar campo' al abrir en Word."
        )

    async def create_docx_version(self, md_path: Optional[str]) -> None:
        if md_path is None:
            logger.error("No path provided for DOCX conversion")
            return

        docx_path = md_path.replace(".md", ".docx")

        async def blocking_create_docx() -> None:
            async with aiofiles.open(md_path, "r", encoding="utf-8") as f:
                md_content = await f.read()

            doc = Document()
            self._add_toc(doc)
            doc.add_page_break()

            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, "html.parser")

            for element in soup:
                if isinstance(element, Tag):
                    if element.name in ["h1", "h2", "h3"]:
                        level = int(element.name[1])
                        heading = doc.add_heading(
                            element.get_text().strip(), level=level
                        )
                        heading.style = f"Heading {level}"
                    elif element.name == "p":
                        doc.add_paragraph(element.get_text().strip())
                    elif element.name == "ul":
                        for item in element.find_all("li", recursive=False):
                            doc.add_paragraph(
                                item.get_text().strip(), style="List Bullet"
                            )
                    elif element.name == "ol":
                        counter = 1
                        for item in element.find_all("li", recursive=False):
                            text = item.get_text().strip()
                            doc.add_paragraph(f"{counter}. {text}", style="List Number")
                            counter += 1

            doc.save(docx_path)

        await blocking_create_docx()

        logger.info(f"DOCX version created: {docx_path}")
