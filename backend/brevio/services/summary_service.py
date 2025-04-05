import asyncio
import json
import logging
import os
import time
from typing import Any, Awaitable, Callable, List, Optional, Tuple

import markdown
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
from openai import AsyncOpenAI
from openai.types.chat import ChatCompletionMessageParam

from brevio.constants.summary_messages import SummaryMessages
from brevio.enums.role import RoleType
from brevio.models.file_config_model import FileConfig
from brevio.models.prompt_config_model import PromptConfig
from brevio.models.response_model import SummaryResponse

from .advanced_content_generator import AdvancedContentGenerator

load_dotenv()

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
logger.propagate = False
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)


class SummaryService:
    def __init__(self) -> None:
        logger.info("Initializing SummaryService for a new user.")
        self.max_tokens = int(os.getenv("MAX_TOKENS") or 4096)
        self.max_tokens_per_chunk = int(os.getenv("MAX_TOKENS_PER_CHUNK") or 3000)
        self.tokens_per_minute = int(os.getenv("TOKENS_PER_MINUTE") or 200000)
        self.temperature = float(os.getenv("TEMPERATURE") or 0.2)
        self.chunk_overlap = 100
        self.max_concurrent_chunks = 10
        self.max_concurrent_files = 5
        self.max_concurrent_requests = 10
        self.token_bucket = self.tokens_per_minute
        self.last_token_reset = time.time()
        self.task_queue: asyncio.Queue[
            Tuple[Callable[..., Awaitable[Any]], List[Any]]
        ] = asyncio.Queue()
        self.running = False
        api_key = os.getenv("DEEPSEEK_API_KEY")

        if not api_key:
            logger.error("DEEPSEEK_API_KEY is not set.")
            raise ValueError("API key not configured")

        self.client = AsyncOpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        self.semaphore = asyncio.Semaphore(self.max_concurrent_requests)
        self.file_semaphore = asyncio.Semaphore(self.max_concurrent_files)
        self.generator = AdvancedContentGenerator()

    async def start(self) -> None:
        if not self.running:
            self.running = True
            asyncio.create_task(self._process_queue())
            logger.info("SummaryService queue processor started.")

    def chunk_text(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunks.append(text[start:end])
            start = end - overlap if end - overlap > start else end
        logger.debug(
            f"Text split into {len(chunks)} chunks with size {chunk_size} and overlap {overlap}."
        )
        return chunks

    async def _update_token_bucket(self) -> None:
        current_time = time.time()
        elapsed_minutes = (current_time - self.last_token_reset) / 60
        if elapsed_minutes >= 1:
            self.token_bucket = min(
                self.tokens_per_minute,
                self.token_bucket + int(self.tokens_per_minute * elapsed_minutes),
            )
            self.last_token_reset = current_time
            logger.info(f"Token bucket updated. Current tokens: {self.token_bucket}")

    async def _check_token_limit(self, tokens_needed: int) -> bool:
        await self._update_token_bucket()
        if self.token_bucket >= tokens_needed:
            return True
        logger.warning(
            f"Token limit reached. Tokens needed: {tokens_needed}, available: {self.token_bucket}. Queuing task."
        )
        return False

    async def generate_summary_chunk(
        self, index: int, chunk: str, prompt: str, accumulated_summary: str
    ) -> Tuple[int, str, int]:
        try:
            truncated_summary = (
                accumulated_summary[-2000:]
                if len(accumulated_summary) > 2000
                else accumulated_summary
            )
            full_prompt = (
                f"{prompt}\n\n"
                f"Contexto previo del resumen: {truncated_summary}\n"
                f"Instrucciones: Proporciona un resumen detallado del siguiente texto, integrando nueva información con el contexto previo de manera coherente. "
                f"Incluye ejemplos, explicaciones y cualquier detalle que facilite el estudio del tema. "
                f"Organiza el resumen en secciones o puntos clave para una mejor comprensión."
            )

            messages: List[ChatCompletionMessageParam] = [
                {"role": RoleType.SYSTEM.value, "content": full_prompt},
                {"role": RoleType.USER.value, "content": chunk},
            ]

            async with self.semaphore:
                logger.debug(
                    f"Sending request for chunk {index}. Chunk preview: {chunk[:500]}..."
                )
                response = await self.client.chat.completions.create(
                    model="deepseek-chat",
                    messages=messages,
                    max_tokens=self.max_tokens,
                    temperature=self.temperature,
                )

                if not hasattr(response, "choices") or not response.choices:
                    logger.error(f"Invalid response for chunk {index}: {response}")
                    return (
                        index,
                        f"Error procesando chunk {index}: Respuesta inválida de la API",
                        0,
                    )

                content = response.choices[0].message.content
                if content is None:
                    logger.error(f"Chunk {index} has no content in response.")
                    return index, "Error: No content in response", 0

                summary = content.strip()
                tokens_used = (
                    response.usage.total_tokens if response.usage is not None else 0
                )
                self.token_bucket -= tokens_used
                logger.info(
                    f"Chunk {index} processed. Tokens used: {tokens_used}, Remaining: {self.token_bucket}"
                )
                return index, summary, tokens_used
        except Exception as e:
            logger.error(f"Error processing chunk {index}: {str(e)}")
            return index, f"Error procesando chunk", 0

    async def process_chunks_in_groups(
        self, chunks: List[str], prompt: str
    ) -> Tuple[str, int]:
        accumulated_summary = ""
        total_tokens_used = 0
        chunk_summaries: List[Optional[str]] = [None] * len(chunks)
        group_size = 5

        for group_start in range(0, len(chunks), group_size):
            group_end = min(group_start + group_size, len(chunks))
            group_chunks = chunks[group_start:group_end]
            group_indices = list(range(group_start, group_end))

            logger.info(f"Processing chunk group {group_start} to {group_end - 1}")

            total_tokens_needed = sum(len(chunk) // 4 + 500 for chunk in group_chunks)
            logger.info(f"Tokens needed for this group: {total_tokens_needed}")

            while not await self._check_token_limit(total_tokens_needed):
                logger.warning(
                    f"Token limit reached. Tokens needed: {total_tokens_needed}, "
                    f"available: {self.token_bucket}. Waiting 5 seconds for token regeneration."
                )
                await asyncio.sleep(5)
                await self._update_token_bucket()

            tasks = [
                self.generate_summary_chunk(index, chunk, prompt, accumulated_summary)
                for index, chunk in zip(group_indices, group_chunks)
            ]
            results = await asyncio.gather(*tasks, return_exceptions=True)

            for result in results:
                if isinstance(result, tuple):
                    index, chunk_summary, tokens_used = result
                    chunk_summaries[index] = chunk_summary
                    total_tokens_used += tokens_used
                    accumulated_summary += (
                        "\n" + chunk_summary if accumulated_summary else chunk_summary
                    )
                else:
                    logger.warning(f"Task failed for chunk: {str(result)}")

        full_summary = "\n".join(
            summary for summary in chunk_summaries if summary is not None
        ).strip()
        return full_summary, total_tokens_used

    async def _process_queue(self) -> None:
        while self.running:
            if self.task_queue.empty():
                await asyncio.sleep(1)
                continue
            func, args = await self.task_queue.get()
            try:
                logger.info("Processing a queued task.")
                await func(*args)
            except Exception as e:
                logger.error(f"Error processing queued task: {str(e)}")
            finally:
                self.task_queue.task_done()
                logger.info("Task completed and removed from queue.")

    async def generate_summary_documents(
        self, prompt_config: PromptConfig, file_configs: List[FileConfig]
    ) -> List[SummaryResponse]:
        logger.info(f"Starting summary generation for {len(file_configs)} documents.")
        try:
            prompt = await self.generator.generate_prompt(
                category=prompt_config.category,
                style=prompt_config.style,
                output_format=prompt_config.format.value,
                lang=prompt_config.language,
                source_type=prompt_config.source_types,
            )

            tasks = [
                self._process_single_document(prompt, file_config)
                for file_config in file_configs
            ]
            results = await asyncio.gather(*tasks)

            logger.info("All document summaries completed.")
            return results
        except Exception as e:
            logger.error(
                f"Error during multiple document summary generation: {str(e)}",
                exc_info=True,
            )
            raise

    async def _process_single_document(
        self, prompt: str, file_config: FileConfig
    ) -> SummaryResponse:
        from brevio.managers.directory_manager import DirectoryManager

        directory_manager = DirectoryManager()
        logger.info(f"Processing document: {file_config.document_path}")
        async with self.file_semaphore:
            try:
                # Ensure required paths are provided
                assert (
                    file_config.document_path is not None
                ), "document_path must be provided"
                assert (
                    file_config.summary_path is not None
                ), "summary_path must be provided"

                directory_manager.validate_paths(file_config.document_path)
                file_extension = (
                    file_config.document_path.split(".")[-1].lower()
                    if file_config.document_path
                    else ""
                )
                logger.info(f"Detected file extension: {file_extension}")

                if file_extension == "pdf":
                    fragments = list(
                        directory_manager.read_pdf(file_config.document_path)
                    )
                    full_text = "\n".join(fragments)
                elif file_extension == "docx":
                    doc = Document(file_config.document_path)
                    full_text = "\n".join(
                        [para.text for para in doc.paragraphs if para.text.strip()]
                    )
                else:
                    raise ValueError(f"Unsupported file type: {file_extension}")

                chars_per_token = 4
                chunk_size = self.max_tokens_per_chunk * chars_per_token
                overlap = self.chunk_overlap

                chunks = self.chunk_text(full_text, chunk_size, overlap)
                logger.info(
                    f"Document split into {len(chunks)} chunks of ~{self.max_tokens_per_chunk} tokens each."
                )

                full_summary, total_tokens_used = await self.process_chunks_in_groups(
                    chunks, prompt
                )
                final_summary = await self.postprocess_summary(full_summary)
                directory_manager.write_summary(final_summary, file_config.summary_path)
                self._create_docx_version(file_config.summary_path)

                logger.info(
                    f"Summary for {file_config.document_path} completed. Total tokens used: {total_tokens_used}"
                )
                return SummaryResponse(
                    success=True,
                    summary=final_summary,
                    message=SummaryMessages.WRITE_SUMMARY.format(
                        file_config.summary_path
                    ),
                )
            except Exception as e:
                logger.error(
                    f"Error processing document {file_config.document_path}: {str(e)}",
                    exc_info=True,
                )
                raise

    async def generate_summary(
        self,
        prompt_config: PromptConfig,
        file_configs: Optional[List[FileConfig]] = None,
        file_config: Optional[FileConfig] = None,
    ) -> List[SummaryResponse]:
        logger.info("Starting summary generation.")
        if file_config is not None and file_configs is None:
            file_configs = [file_config]
        elif file_configs is None:
            raise ValueError("Debe proporcionar file_configs o file_config.")
        try:
            prompt = await self.generator.generate_prompt(
                category=prompt_config.category,
                style=prompt_config.style,
                output_format=prompt_config.format.value,
                lang=prompt_config.language,
                source_type=prompt_config.source_types,
            )

            tasks = [
                self._process_single_transcription(prompt, fc) for fc in file_configs
            ]
            results = await asyncio.gather(*tasks)

            logger.info("All summaries completed.")
            return results
        except Exception as e:
            logger.error(f"Error during summary generation: {str(e)}", exc_info=True)
            raise

    async def _process_single_transcription(
        self, prompt: str, file_config: FileConfig
    ) -> SummaryResponse:
        from brevio.managers.directory_manager import DirectoryManager

        directory_manager = DirectoryManager()
        logger.info(f"Processing transcription: {file_config.transcription_path}")
        async with self.file_semaphore:
            try:
                # Ensure required paths are provided
                assert (
                    file_config.transcription_path is not None
                ), "transcription_path must be provided"
                assert (
                    file_config.summary_path is not None
                ), "summary_path must be provided"

                directory_manager.validate_paths(file_config.transcription_path)
                transcription = directory_manager.read_transcription(
                    file_config.transcription_path
                )
                logger.info("Transcription read successfully.")

                tokens_per_chunk = 3500
                chars_per_token = 4
                chunk_size = tokens_per_chunk * chars_per_token
                overlap = self.chunk_overlap

                chunks = self.chunk_text(transcription, chunk_size, overlap)
                logger.info(
                    f"Text split into {len(chunks)} chunks of ~{tokens_per_chunk} tokens each."
                )

                full_summary, total_tokens_used = await self.process_chunks_in_groups(
                    chunks, prompt
                )
                final_summary = await self.postprocess_summary(full_summary)

                directory_manager.write_summary(final_summary, file_config.summary_path)
                self._create_docx_version(file_config.summary_path)

                logger.info(
                    f"Summary for {file_config.transcription_path} completed. Total tokens used: {total_tokens_used}"
                )
                return SummaryResponse(
                    success=True,
                    summary=final_summary,
                    message=SummaryMessages.WRITE_SUMMARY.format(
                        file_config.summary_path
                    ),
                )
            except Exception as e:
                logger.error(
                    f"Error processing transcription {file_config.transcription_path}: {str(e)}",
                    exc_info=True,
                )
                raise

    async def postprocess_summary(self, summary: str) -> str:
        logger.info("Postprocessing summary to remove redundancies.")
        try:
            tokens_needed = len(summary) // 4 + 500
            if tokens_needed <= self.max_tokens:
                if not await self._check_token_limit(tokens_needed):
                    await self.task_queue.put((self.postprocess_summary, [summary]))
                    logger.info("Postprocessing task queued due to token limit.")
                    return summary
                response = await self.client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {
                            "role": "system",
                            "content": "Eres un editor experto en optimizar textos eliminando redundancias.",
                        },
                        {
                            "role": "user",
                            "content": (
                                "Revisa el siguiente resumen y elimina exclusivamente redundancias o información repetida, como textos, frases o ideas duplicadas. "
                                "No simplifiques, no resumas ni reduzcas el contenido de ninguna manera; conserva intactos todos los detalles, datos y elementos importantes. "
                                "Asegúrate de que el texto final sea claro, cohesivo y bien organizado, sin alterar su estructura ni su significado original."
                                f"\n\n{summary}"
                            ),
                        },
                    ],
                    max_tokens=self.max_tokens,
                    temperature=self.temperature,
                )
                if response.usage is not None:
                    self.token_bucket -= response.usage.total_tokens
                    logger.info(
                        f"Postprocessing completed. Tokens used: {response.usage.total_tokens}"
                    )
                else:
                    logger.warning("No usage information in response.")
                    logger.info("Postprocessing completed. Tokens used: 0")

                content = response.choices[0].message.content
                return content.strip() if content is not None else ""

            chunk_size = (self.max_tokens - 500) * 4
            chunks = self.chunk_text(summary, chunk_size, self.chunk_overlap)
            logger.info(f"Summary split into {len(chunks)} chunks for postprocessing.")

            final_summary, total_tokens_used = await self.process_chunks_in_groups(
                chunks,
                "Eres un editor experto en mejorar textos. "
                "Revisa el siguiente texto y elimina únicamente las redundancias o información repetida, "
                "asegurando que el texto sea claro, cohesivo y bien organizado. "
                "No resumas ni reduzcas el contenido adicionalmente; mantén todos los detalles importantes intactos.",
            )
            logger.info(
                f"Postprocessing of chunks completed. Total tokens used: {total_tokens_used}"
            )
            return final_summary
        except Exception as e:
            logger.error(f"Postprocessing failed: {str(e)}")
            raise

    def _add_toc(self, doc: DocumentType) -> None:
        doc.add_heading("Índice", level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char)
        instr_text = OxmlElement("w:instrText")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr_text)
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char)
        doc.add_paragraph(
            "Nota: Haga clic derecho en el índice y seleccione 'Actualizar campo' al abrir en Word."
        )

    def _create_docx_version(self, md_path: Optional[str]) -> None:
        if md_path is None:
            logger.error("No path provided for DOCX conversion")
            return

        docx_path = md_path.replace(".md", ".docx")
        try:
            with open(md_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            doc = Document()
            self._add_toc(doc)
            doc.add_page_break()

            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, "html.parser")

            for element in soup:
                if isinstance(element, Tag):
                    if element.name in ["h1", "h2", "h3"]:
                        level = int(element.name[1])
                        heading = doc.add_heading(
                            element.get_text().strip(), level=level
                        )
                        heading.style = f"Heading {level}"
                    elif element.name == "p":
                        doc.add_paragraph(element.get_text().strip())
                    elif element.name == "ul":
                        for item in element.find_all("li", recursive=False):  # type: ignore
                            doc.add_paragraph(
                                item.get_text().strip(), style="List Bullet"
                            )
                    elif element.name == "ol":
                        counter = 1
                        for item in element.find_all("li", recursive=False):  # type: ignore
                            text = item.get_text().strip()
                            doc.add_paragraph(f"{counter}. {text}", style="List Number")
                            counter += 1

            doc.save(docx_path)
            logger.info(f"DOCX version created: {docx_path}")
        except Exception as e:
            logger.error(f"DOCX conversion failed: {str(e)}")
            raise

    async def generate_summary_document(
        self, prompt_config: PromptConfig, file_config: FileConfig
    ) -> SummaryResponse:
        logger.info(
            f"Starting summary generation for single document: {file_config.document_path}"
        )
        try:
            results = await self.generate_summary_documents(
                prompt_config, [file_config]
            )
            if not results:
                raise ValueError("No results returned from summary generation")
            return results[0]
        except Exception as e:
            logger.error(f"Error generating summary for document: {str(e)}")
            raise


async def process_user_summaries(
    user_id: int, prompt_config: PromptConfig, file_configs: List[FileConfig]
) -> None:
    service = SummaryService()
    await service.start()
    logger.info(f"User {user_id} starting summary generation.")
    try:
        results = await service.generate_summary_documents(prompt_config, file_configs)
        for result in results:
            logger.info(f"User {user_id} summary: {result.summary[:50]}...")
    except Exception as e:
        logger.error(f"Error processing summaries for user {user_id}: {str(e)}")
        raise
